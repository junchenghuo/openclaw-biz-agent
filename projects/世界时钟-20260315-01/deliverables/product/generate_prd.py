#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# Title
title = doc.add_heading('群流程回归-世界时钟-20260315-01', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('产品需求说明书（PRD）')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.italic = True
subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)

# 1. 项目概述
doc.add_heading('1. 项目概述', level=1)
doc.add_paragraph('项目名称：群流程回归-世界时钟-20260315-01')
doc.add_paragraph('项目目标：展示8国时钟，支持实时时间显示与时区查询')
doc.add_paragraph('项目范围：前端 + 后端 + 测试 + 部署')
doc.add_paragraph('验收标准：8国时区可查询且误差<1秒')

# 2. 目标用户
doc.add_heading('2. 目标用户', level=1)
doc.add_paragraph('• 需要跨时区协作的团队成员')
doc.add_paragraph('• 关注全球时区的个人用户')

# 3. 功能需求
doc.add_heading('3. 功能需求', level=1)

doc.add_heading('3.1 时钟展示', level=2)
doc.add_paragraph('• 展示8个国家的当前时间（建议：中国、美国、英国、日本、德国、澳大利亚、法国、俄罗斯）')
doc.add_paragraph('• 时间自动实时更新，误差<1秒')
doc.add_paragraph('• 显示各国时区名称和UTC偏移量')

doc.add_heading('3.2 时区查询', level=2)
doc.add_paragraph('• 用户可选择特定国家/时区查看详情')
doc.add_paragraph('• 支持时间换算功能')

doc.add_heading('3.3 界面交互', level=2)
doc.add_paragraph('• 清晰的时钟卡片布局')
doc.add_paragraph('• 响应式设计，支持移动端和桌面端')

# 4. 非功能需求
doc.add_heading('4. 非功能需求', level=1)
doc.add_paragraph('• 性能：页面加载时间<2秒')
doc.add_paragraph('• 准确性：时间误差<1秒')
doc.add_paragraph('• 可用性：支持主流浏览器（Chrome、Firefox、Safari、Edge）')
doc.add_paragraph('• 可访问性：支持基本的屏幕阅读器')

# 5. 技术架构
doc.add_heading('5. 技术架构', level=1)
doc.add_paragraph('• 前端：Web技术实现响应式界面')
doc.add_paragraph('• 后端：提供标准时间API接口')
doc.add_paragraph('• 部署：支持容器化部署')

# 6. 验收标准
doc.add_heading('6. 验收标准', level=1)
doc.add_paragraph('✓ 成功展示8个国家的实时时钟')
doc.add_paragraph('✓ 时间误差控制在1秒以内')
doc.add_paragraph('✓ 界面在移动端和桌面端均正常显示')
doc.add_paragraph('✓ 支持至少一种时区查询/换算功能')
doc.add_paragraph('✓ 完成部署并可公网访问')

# 7. 里程碑
doc.add_heading('7. 里程碑', level=1)
doc.add_paragraph('• M1：PRD + 原型图评审通过（Day 1）')
doc.add_paragraph('• M2：架构设计完成（Day 2）')
doc.add_paragraph('• M3：开发完成 + 联调通过（Day 4）')
doc.add_paragraph('• M4：测试通过 + 部署上线（Day 5）')

# 8. 风险与依赖
doc.add_heading('8. 风险与依赖', level=1)
doc.add_paragraph('• 风险：时间同步API的稳定性')
doc.add_paragraph('• 依赖：后端时间服务接口')

# Save
output_path = '/Users/imac/midCreate/openclaw-workspaces/ai-team/projects/世界时钟-20260315-01/deliverables/product/世界时钟-20260315-01-产品需求说明书.doc'
doc.save(output_path)
print(f"PRD created: {output_path}")
